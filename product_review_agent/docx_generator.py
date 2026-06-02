# -*- coding: utf-8 -*-
"""
审核报告 Word 文档生成器

共享模块，供 Web API 和飞书 Bot 共同使用。
生成格式化的 .docx 审核报告文档。

报告结构：
  一、立项信息
  二、市场与定价
  三、人群与场景分析
  四、专项分析（含逐模块差距矩阵表格）
  五、综合评估
  附：设计要求
"""

from __future__ import annotations

import json
import tempfile
import os
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ============================================================
# 颜色常量
# ============================================================
COLOR_DARK = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_GREEN = RGBColor(0x00, 0x80, 0x00)
COLOR_RED = RGBColor(0xCC, 0x33, 0x00)
COLOR_BLUE = RGBColor(0x00, 0x78, 0xD4)
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)

GAP_COLORS = {
    "高": RGBColor(0xCC, 0x33, 0x00),
    "中": RGBColor(0xFF, 0x99, 0x00),
    "低": RGBColor(0x00, 0x80, 0x00),
    "无": RGBColor(0x99, 0x99, 0x99),
}


# ============================================================
# 主入口
# ============================================================

def generate_review_docx(
    file_name: str = "report",
    task_label: str = "",
    overall_score: int = 0,
    risk_level: str = "未知",
    project_data: Optional[dict] = None,
    specific_score: Optional[dict] = None,
    specific_analysis: Optional[dict] = None,
    common_scores: Optional[dict] = None,
    report_text: str = "",
    save_dir: Optional[str] = None,
) -> str:
    """
    生成审核报告 Word 文档。

    Args:
        file_name: 原始文件名
        task_label: 审核类型标签
        overall_score: 综合评分
        risk_level: 风险等级
        project_data: Excel 解析出的项目数据
        specific_score: 专项分析评分（dict 或 AnalyzerScore）
        specific_analysis: 专项分析原始结果（含 module_comparison 等）
        common_scores: 公共分析评分（含 audience_scenario）
        report_text: 完整审核报告文本（备用，提取设计要求等）
        save_dir: 保存目录

    Returns:
        生成的 .docx 文件绝对路径
    """
    project_data = project_data or {}
    specific_score = specific_score or {}
    specific_analysis = specific_analysis or {}
    common_scores = common_scores or {}

    doc = Document()

    # ---- 全局样式 ----
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3

    # ---- 标题 ----
    title = doc.add_heading("产品立项审核报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = COLOR_DARK

    # ---- 评分总览条 ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"综合评分: {overall_score}/100  |  风险等级: {risk_level}  |  {task_label}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = _score_color(overall_score)

    # ================================================================
    # 一、立项信息
    # ================================================================
    doc.add_heading("一、立项信息", level=1)

    info_items = [
        ("产品名称", project_data.get("project_name") or project_data.get("product_name", "(未填写)")),
        ("品牌", project_data.get("brand", "(未填写)")),
        ("品类", " > ".join(filter(None, [
            project_data.get("categoryl1", "") or project_data.get("category_l1", ""),
            project_data.get("categoryl2", "") or project_data.get("category_l2", ""),
            project_data.get("categoryl3", "") or project_data.get("category_l3", ""),
        ])) or "(未填写)"),
        ("负责人", project_data.get("applicant", "(未填写)")),
        ("立项时间", project_data.get("project_time", "(未填写)")),
        ("设计时间", project_data.get("design_time", "(未填写)")),
        ("打样时间", project_data.get("proofing_time", "(未填写)")),
        ("上架时间", project_data.get("launch_time", "(未填写)")),
        ("审核类型", task_label),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]
    _add_kv_table(doc, info_items)

    # ================================================================
    # 二、市场与定价
    # ================================================================
    doc.add_heading("二、市场与定价", level=1)

    market_items = [
        ("市场规模", project_data.get("market_size", "(未填写)")),
        ("目标销量", project_data.get("estimated_sales", "(未填写)")),
        ("定价", project_data.get("pricing", "(未填写)")),
        ("毛利率", project_data.get("gfm", "(未填写)")),
        ("ERP成本", project_data.get("ERP_price", "(未填写)")),
    ]
    _add_kv_table(doc, market_items)

    # 竞品对比信息
    comparison = project_data.get("product_comparison")
    if not comparison and project_data.get("competitor_price"):
        comparison = {"competitor_price": project_data["competitor_price"]}
    if comparison:
        doc.add_paragraph("")  # 空行
        p = doc.add_paragraph()
        run = p.add_run("竞品对标")
        run.bold = True
        run.font.size = Pt(11)

        if isinstance(comparison, list):
            for i, comp in enumerate(comparison):
                prefix = f"竞品{i + 1}" if len(comparison) > 1 else "竞品"
                if isinstance(comp, dict):
                    comp_items = [
                        (f"{prefix} - 对手商品", comp.get("comparison_name", "(未填写)")),
                        (f"{prefix} - 对手卖点(复制)", comp.get("selling_point", "(未填写)")),
                        (f"{prefix} - 对手卖点(超越)", comp.get("improving_point", "(未填写)")),
                    ]
                    _add_kv_table(doc, comp_items)
                else:
                    doc.add_paragraph(f"  {prefix}: {comp}")
        elif isinstance(comparison, dict):
            comp_items = [
                ("对手商品", comparison.get("comparison_name", "(未填写)")),
                ("对手卖点(复制)", comparison.get("selling_point", "(未填写)")),
                ("对手卖点(超越)", comparison.get("improving_point", "(未填写)")),
            ]
            _add_kv_table(doc, comp_items)

    # ================================================================
    # 三、人群与场景分析
    # ================================================================
    doc.add_heading("三、人群与场景分析", level=1)

    as_score = common_scores.get("audience_scenario", {})
    if as_score:
        _add_audience_scenario_section(doc, as_score)
    else:
        # 兼容旧格式（单独的 audience + scenario）
        audience = common_scores.get("audience", {})
        scenario = common_scores.get("scenario", {})
        if audience:
            _add_analysis_only_section(doc, audience, "人群")
        if scenario:
            _add_analysis_only_section(doc, scenario, "场景")
        if not audience and not scenario:
            doc.add_paragraph("(暂无人群与场景分析数据)")

    # ================================================================
    # 四、专项分析
    # ================================================================
    doc.add_heading(f"四、{task_label or '专项'}分析", level=1)

    if specific_score:
        _add_score_section(doc, specific_score, "专项")

        # 爆品升级5维度详细评语（每个维度单独段落展示）
        if task_label and "爆品" in task_label:
            _add_hot_upgrade_dimension_details(doc, specific_score)

        # 品类缺失维度详细评语
        if task_label and "品类" in task_label:
            scenario = specific_analysis.get("scenario", "B") if specific_analysis else "B"
            _add_category_gap_dimension_details(doc, specific_score, scenario)

        # 竞品升级5维度详细评语
        if task_label and "竞品" in task_label:
            _add_competitor_upgrade_dimension_details(doc, specific_score)

        # 未起量迭代5维度详细评语
        if task_label and "未起量" in task_label:
            _add_low_sale_iterate_dimension_details(doc, specific_score)

    # 立项信息详情
    proj = specific_analysis.get("project_data", {}) if specific_analysis else {}
    if proj:
        detail_items = []
        # 爆品升级字段
        if proj.get("upgrade_modules"):
            detail_items.append(("升级模块", proj["upgrade_modules"]))
        if proj.get("design_purpose"):
            detail_items.append(("设计目的", proj["design_purpose"]))
        if proj.get("product_hotpoint"):
            detail_items.append(("原有卖点", proj["product_hotpoint"]))
        if proj.get("upgrade_valiable"):
            detail_items.append(("可行性说明", proj["upgrade_valiable"]))
        # 竞品升级字段
        if proj.get("competitor_strengths_copy"):
            detail_items.append(("竞品卖点-需复制", proj["competitor_strengths_copy"]))
        if proj.get("competitor_advantage"):
            detail_items.append(("竞品卖点-需超越", proj["competitor_advantage"]))
        # 未起量迭代字段
        if proj.get("failure_analysis"):
            detail_items.append(("没卖好的原因", proj["failure_analysis"]))
        if proj.get("current_issues"):
            detail_items.append(("当前产品问题", proj["current_issues"]))
        if proj.get("sales_data_desc"):
            detail_items.append(("销量现状", proj["sales_data_desc"]))
        # 品类缺失字段
        design_content = proj.get("design_content") or proj.get("upgrade_modules", "")
        if design_content and not proj.get("upgrade_modules"):
            detail_items.append(("设计内容", design_content))
        feasibility = proj.get("feasibility_analysis") or proj.get("upgrade_valiable", "")
        if feasibility and not proj.get("upgrade_valiable"):
            detail_items.append(("可行性分析", feasibility))
        if proj.get("audience_consistent"):
            detail_items.append(("二级品类人群一致性", proj["audience_consistent"]))
        if proj.get("market_size"):
            detail_items.append(("市场大小", proj["market_size"]))
        if proj.get("estimated_sales"):
            detail_items.append(("目标销售额", proj["estimated_sales"]))
        # 通用字段
        if proj.get("pricing"):
            detail_items.append(("定价", proj["pricing"]))
        if proj.get("erp_cost"):
            detail_items.append(("ERP成本", str(proj["erp_cost"])))
        if proj.get("competitor_price"):
            detail_items.append(("竞品价格", str(proj["competitor_price"])))
        if detail_items:
            _add_kv_table(doc, detail_items)

    # 品类缺失场景判断
    if task_label and "品类" in task_label and specific_analysis:
        _add_category_gap_scenario_section(doc, specific_analysis)

    # CBB匹配详情
    if specific_analysis:
        _add_cbb_match_section(doc, specific_analysis)

    # 品类市场概况
    if task_label and "品类" in task_label and specific_analysis:
        _add_market_overview_section(doc, specific_analysis)

    # VL对比拆解详细内容
    vl_report = specific_analysis.get("vl_report", {}) if specific_analysis else {}
    if vl_report and "error" not in vl_report:
        # 对比模式（有section3）
        if vl_report.get("section3_module_comparison"):
            _add_vl_report_sections(doc, vl_report)
        # 单拆模式（场景B：只有b_level）
        elif vl_report.get("section2_abc_modules", {}).get("b_level"):
            _add_single_vl_report(doc, vl_report)

    # 模块CBB分类映射
    llm_scoring = specific_analysis.get("llm_scoring", {}) if specific_analysis else {}
    mr = llm_scoring.get("module_reuse", {})
    if mr and not mr.get("_error"):
        mapping = mr.get("vl_to_cbb_mapping", [])
        if mapping:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("VL模块→CBB分类映射")
            run.bold = True
            run.font.size = Pt(11)
            for m in mapping[:10]:
                doc.add_paragraph(
                    f"{m.get('vl_module', '?')} → {m.get('cbb_category', '?')} ({m.get('side', '')})",
                    style="List Bullet"
                )
        core_cats = mr.get("core_categories", [])
        if core_cats:
            doc.add_paragraph(f"核心CBB分类: {', '.join(core_cats)}")
        missing_cats = mr.get("missing_categories", [])
        if missing_cats:
            p = doc.add_paragraph()
            run = p.add_run(f"缺失分类: {', '.join(missing_cats)}")
            run.font.color.rgb = COLOR_RED
        reuse_cats = mr.get("reuse_categories", [])
        if reuse_cats:
            doc.add_paragraph(f"可复用分类: {', '.join(reuse_cats)}")

    # 逐模块差距矩阵（Word 表格）
    module_comparison = _extract_module_comparison(specific_analysis)
    if module_comparison:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("逐模块差距矩阵")
        run.bold = True
        run.font.size = Pt(11)

        _add_module_comparison_table(doc, module_comparison)

    # 升级路线图
    upgrade_roadmap = _extract_roadmap(specific_analysis)
    if upgrade_roadmap:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("升级路线图")
        run.bold = True
        run.font.size = Pt(11)

        _add_roadmap_table(doc, upgrade_roadmap)

    # 分析摘要
    summary = _extract_summary(specific_analysis)
    if summary:
        _add_summary_section(doc, summary)

    if not specific_score and not module_comparison and not vl_report:
        doc.add_paragraph("(暂无专项分析数据)")

    # ================================================================
    # 五、综合评估
    # ================================================================
    doc.add_heading("五、综合评估", level=1)

    # 综合评分
    p = doc.add_paragraph()
    run = p.add_run(f"综合评分: {overall_score}/100")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = _score_color(overall_score)

    p = doc.add_paragraph()
    run = p.add_run(f"风险等级: {risk_level}")
    run.bold = True
    if risk_level == "高":
        run.font.color.rgb = COLOR_RED
    elif risk_level == "中":
        run.font.color.rgb = RGBColor(0xFF, 0x99, 0x00)
    else:
        run.font.color.rgb = COLOR_GREEN

    # 专项评分
    sp_total = _get_specific_total(specific_score)
    p = doc.add_paragraph()
    p.add_run("专项评分: ").bold = True
    p.add_run(f"{sp_total}/100（人群与场景仅做分析参考，不纳入评分）")

    # 各部分优劣势汇总
    all_strengths = []
    all_weaknesses = []
    all_suggestions = []

    if as_score:
        all_strengths.extend(as_score.get("strengths", []))
        all_weaknesses.extend(as_score.get("weaknesses", []))
        all_suggestions.extend(as_score.get("suggestions", []))

    if specific_score:
        all_strengths.extend(specific_score.get("strengths", []))
        all_weaknesses.extend(specific_score.get("weaknesses", []))
        all_suggestions.extend(specific_score.get("suggestions", []))

    if all_strengths:
        _add_bullet_list(doc, "优势", all_strengths, COLOR_GREEN, "✅")
    if all_weaknesses:
        _add_bullet_list(doc, "不足", all_weaknesses, COLOR_RED, "⚠️")
    if all_suggestions:
        _add_bullet_list(doc, "改进建议", all_suggestions, COLOR_BLUE, "💡")

    # ================================================================
    # 附：设计要求
    # ================================================================
    design_req = project_data.get("design_requirements", "")
    if not design_req:
        design_req = project_data.get("design_requirement", "")
    if not design_req:
        design_req = project_data.get("design_require", "")
    if not design_req and project_data.get("design_purpose"):
        design_req = project_data["design_purpose"]
    if not design_req:
        # 从 report_text 中提取
        design_req = _extract_design_requirements(report_text)

    if design_req:
        doc.add_heading("附：设计要求", level=1)
        if isinstance(design_req, list):
            for item in design_req:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            for line in str(design_req).split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

    # ---- 保存文件 ----
    if save_dir is None:
        save_dir = tempfile.mkdtemp()
    base_name = file_name.replace(".xlsx", "").replace(".xls", "")
    docx_path = os.path.join(save_dir, f"审核报告_{base_name}.docx")
    doc.save(docx_path)

    return docx_path


# ============================================================
# 内部辅助函数
# ============================================================

def _score_color(score: int) -> RGBColor:
    """根据分值返回颜色"""
    if score >= 75:
        return COLOR_GREEN
    elif score >= 50:
        return RGBColor(0xFF, 0x99, 0x00)
    else:
        return COLOR_RED


def _add_kv_table(doc, items: list[tuple[str, str]]):
    """添加键值对表格"""
    table = doc.add_table(rows=len(items), cols=2, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(items):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "微软雅黑"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        # 标签列加粗
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True


def _add_score_section(doc, score_data: dict, section_name: str):
    """向文档添加一个评分板块（评分明细+优势+不足+建议）"""
    total = score_data.get("total_score", 0)
    p = doc.add_paragraph()
    run = p.add_run(f"综合评分: {total}/100")
    run.bold = True
    run.font.size = Pt(11)

    # 评分维度表
    dimensions = score_data.get("dimensions", {})
    if dimensions:
        if isinstance(dimensions, dict):
            dim_items = [(name, info) for name, info in dimensions.items()
                         if isinstance(info, dict)]
        elif isinstance(dimensions, list):
            dim_items = [(d.get("name", ""), d) for d in dimensions
                         if isinstance(d, dict)]
        else:
            dim_items = []
        if dim_items:
            t = doc.add_table(rows=len(dim_items) + 1, cols=3, style="Light Grid Accent 1")
            for j, header in enumerate(["维度", "得分", "评价"]):
                t.rows[0].cells[j].text = header
                for run in t.rows[0].cells[j].paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(9)
            for i, (name, info) in enumerate(dim_items):
                t.rows[i + 1].cells[0].text = name
                max_s = info.get("max_score", 25)
                t.rows[i + 1].cells[1].text = f"{info.get('score', 0)}/{max_s}"
                # 多行reason拆分到单元格
                reason = info.get("reason", "")
                cell = t.rows[i + 1].cells[2]
                # 清除默认段落
                for p in cell.paragraphs:
                    p.clear()
                sub_lines = [r.strip() for r in reason.split("\n") if r.strip()]
                if sub_lines:
                    first = True
                    for sl in sub_lines:
                        if first:
                            cell.paragraphs[0].add_run(sl)
                            first = False
                        else:
                            new_p = cell.add_paragraph()
                            new_p.add_run(sl)
                for cell_item in t.rows[i + 1].cells:
                    for paragraph in cell_item.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.name = "微软雅黑"
                            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 优势/不足/建议
    for label, key, color in [
        ("优势", "strengths", COLOR_GREEN),
        ("不足", "weaknesses", COLOR_RED),
        ("改进建议", "suggestions", COLOR_BLUE),
    ]:
        items = score_data.get(key, [])
        if items:
            _add_bullet_list(doc, label, items, color)


def _add_analysis_only_section(doc, score_data: dict, section_name: str):
    """添加仅分析内容的板块（不显示分数）"""
    analysis = score_data.get("analysis", {})
    if analysis:
        for key, text in analysis.items():
            if text and isinstance(text, str):
                p = doc.add_paragraph()
                run = p.add_run(f"【{key}】")
                run.bold = True
                doc.add_paragraph(text)

    # 维度分析表
    dimensions = score_data.get("dimensions", {})
    if dimensions:
        if isinstance(dimensions, dict):
            dim_items = [(name, info) for name, info in dimensions.items()
                         if isinstance(info, dict)]
        elif isinstance(dimensions, list):
            dim_items = [(d.get("name", ""), d) for d in dimensions
                         if isinstance(d, dict)]
        else:
            dim_items = []
        if dim_items:
            t = doc.add_table(rows=len(dim_items) + 1, cols=2, style="Light Grid Accent 1")
            for j, header in enumerate(["维度", "分析结论"]):
                t.rows[0].cells[j].text = header
                for run in t.rows[0].cells[j].paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(9)
            for i, (name, info) in enumerate(dim_items):
                t.rows[i + 1].cells[0].text = name
                t.rows[i + 1].cells[1].text = info.get("reason", "")
                for cell in t.rows[i + 1].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.name = "微软雅黑"
                            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 优势/不足/建议
    for label, key, color in [
        ("优势", "strengths", COLOR_GREEN),
        ("不足", "weaknesses", COLOR_RED),
        ("改进建议", "suggestions", COLOR_BLUE),
    ]:
        items = score_data.get(key, [])
        if items:
            _add_bullet_list(doc, label, items, color)


def _add_audience_scenario_section(doc, as_score: dict):
    """添加合并版人群与场景分析板块（仅分析，不评分）"""

    # 分析段落
    analysis = as_score.get("analysis", {})
    if analysis:
        dim_labels = {
            "audience_scene_fit": "人群-场景匹配度",
            "insight_depth": "需求洞察深度",
            "data_coverage": "数据支撑与覆盖度",
            "commercial_value": "商业价值判断",
        }
        for key, label in dim_labels.items():
            text = analysis.get(key, "")
            if text:
                p = doc.add_paragraph()
                run = p.add_run(f"【{label}】")
                run.bold = True
                run.font.size = Pt(10.5)
                doc.add_paragraph(text)

    # 专家视角
    expert = as_score.get("expert_analysis", {})
    if expert:
        p = doc.add_paragraph()
        run = p.add_run("专家视角")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_BLUE

        expert_items = []
        if expert.get("target_audience"):
            expert_items.append(("建议目标人群", expert["target_audience"]))
        if expert.get("core_scenarios"):
            expert_items.append(("建议核心场景", expert["core_scenarios"]))
        if expert.get("key_suggestion"):
            expert_items.append(("关键建议", expert["key_suggestion"]))
        if expert_items:
            _add_kv_table(doc, expert_items)

    # 维度分析表（不显示分数）
    scores = as_score.get("scores", {})
    if scores:
        score_items = [(name, info) for name, info in scores.items()
                       if isinstance(info, dict)]
        if score_items:
            t = doc.add_table(rows=len(score_items) + 1, cols=2, style="Light Grid Accent 1")
            for j, header in enumerate(["维度", "分析结论"]):
                t.rows[0].cells[j].text = header
                for run in t.rows[0].cells[j].paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(9)
            for i, (name, info) in enumerate(score_items):
                t.rows[i + 1].cells[0].text = name
                t.rows[i + 1].cells[1].text = info.get("reason", "")
                for cell in t.rows[i + 1].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.name = "微软雅黑"
                            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 优势/不足/建议
    for label, key, color in [
        ("优势", "strengths", COLOR_GREEN),
        ("不足", "weaknesses", COLOR_RED),
        ("改进建议", "suggestions", COLOR_BLUE),
    ]:
        items = as_score.get(key, [])
        if items:
            _add_bullet_list(doc, label, items, color)


def _add_vl_report_sections(doc, vl_report: dict):
    """将VL对比拆解报告的详细内容输出到Word文档"""

    # ── 模块对比分析 (section3) ──
    section3 = vl_report.get("section3_module_comparison", {})
    if section3:
        p = doc.add_paragraph()
        run = p.add_run("── 模块对比分析 ──")
        run.bold = True
        run.font.size = Pt(11)

        reuse_rate = section3.get("overall_reuse_rate", "?")
        reuse_summary = section3.get("reuse_summary", "")
        p = doc.add_paragraph()
        run = p.add_run(f"整体复用率: {reuse_rate}%")
        run.bold = True
        if reuse_summary:
            p.add_run(f"  |  {reuse_summary}")

        same = section3.get("same_modules", [])
        comp_only = section3.get("competitor_only", [])
        own_only = section3.get("own_only", [])
        structural = section3.get("structural_differences", [])

        if same:
            _add_bullet_list(doc, f"相同模块（可复用）: {len(same)}个",
                             [f"{m.get('module_name','?')} | 自家: {str(m.get('own_detail',''))[:40]} | 竞品: {str(m.get('competitor_detail',''))[:40]}"
                              for m in same[:8]], COLOR_GREEN, "✅")
        if comp_only:
            items = []
            for m in comp_only[:8]:
                hit = " ← 升级方向命中" if m.get("upgrade_direction_hit") else ""
                items.append(f"{m.get('module_name','?')} | {str(m.get('detail',''))[:50]}{hit}")
            _add_bullet_list(doc, f"竞品独有（需补齐）: {len(comp_only)}个", items, COLOR_RED, "🔴")
        if own_only:
            items = []
            for m in own_only[:6]:
                adv = " ✓ 优势" if m.get("is_advantage") else ""
                items.append(f"{m.get('module_name','?')} | {str(m.get('detail',''))[:50]}{adv}")
            _add_bullet_list(doc, f"自家独有（差异化）: {len(own_only)}个", items, COLOR_GREEN, "🟢")
        if structural:
            items = []
            for s in structural[:4]:
                items.append(f"{s.get('aspect','?')} | 自家: {str(s.get('own',''))[:30]} | 竞品: {str(s.get('competitor',''))[:30]}")
            _add_bullet_list(doc, f"结构差异: {len(structural)}处", items, COLOR_GRAY, "⚡")

    # ── 升级方向评估 (section4) ──
    dir_score = vl_report.get("section4_upgrade_direction_score", {})
    if dir_score:
        p = doc.add_paragraph()
        run = p.add_run("── 升级方向评估 ──")
        run.bold = True
        run.font.size = Pt(11)

        quality = dir_score.get("direction_quality", "?")
        quality_color = COLOR_GREEN if quality == "精准" else (RGBColor(0xFF, 0x99, 0x00) if quality == "部分命中" else COLOR_RED)
        p = doc.add_paragraph()
        run = p.add_run(f"方向评估: {quality}")
        run.bold = True
        run.font.color.rgb = quality_color

        dir_items = []
        hit_mods = dir_score.get("direction_hit_modules", [])
        miss_mods = dir_score.get("direction_miss_modules", [])
        if hit_mods:
            dir_items.append(("命中模块", ", ".join(str(m) for m in hit_mods[:5])))
        if miss_mods:
            dir_items.append(("未命中模块", ", ".join(str(m) for m in miss_mods[:5])))
        if dir_score.get("reason"):
            dir_items.append(("评估理由", dir_score["reason"]))
        if dir_items:
            _add_kv_table(doc, dir_items)

    # ── 模块复用分析 (section5) ──
    reuse_data = vl_report.get("section5_module_reuse", {})
    if reuse_data:
        p = doc.add_paragraph()
        run = p.add_run("── 模块复用分析 ──")
        run.bold = True
        run.font.size = Pt(11)

        overall_rate = reuse_data.get("overall_reuse_rate", "?")
        core_rate = reuse_data.get("core_module_reuse_rate", "?")
        new_modules = reuse_data.get("new_modules_needed", [])
        reuse_summary = reuse_data.get("reuse_summary", "")

        p = doc.add_paragraph()
        run = p.add_run(f"整体复用率: {overall_rate}%  |  核心模块复用率: {core_rate}%")
        run.bold = True

        if new_modules:
            p = doc.add_paragraph()
            p.add_run(f"需新增模块: {', '.join(str(m) for m in new_modules[:5])}")

        reuse_analysis = reuse_data.get("reuse_analysis", [])
        if reuse_analysis:
            items = [f"{ra.get('module','?')} — 复用度: {ra.get('reusability','?')} | {str(ra.get('reason',''))[:60]}"
                     for ra in reuse_analysis[:6]]
            _add_bullet_list(doc, "逐模块复用评估", items, COLOR_BLUE, "")

        if reuse_summary:
            doc.add_paragraph(f"总结: {reuse_summary}")

    # ── 升级增量价值 (section6) ──
    value_data = vl_report.get("section6_upgrade_value", {})
    if value_data:
        p = doc.add_paragraph()
        run = p.add_run("── 升级增量价值 ──")
        run.bold = True
        run.font.size = Pt(11)

        inc_modules = value_data.get("incremental_modules", [])
        perception = value_data.get("user_perception", "?")
        price_comp = value_data.get("price_competitiveness", "")
        value_summary = value_data.get("value_summary", "")

        val_items = [
            ("新增模块", ", ".join(str(m) for m in inc_modules[:5]) if inc_modules else "无"),
            ("用户感知", perception),
        ]
        if price_comp:
            val_items.append(("价格竞争力", price_comp[:80]))
        _add_kv_table(doc, val_items)

        if value_summary:
            doc.add_paragraph(f"总结: {value_summary}")

    # ── 执行可行性 (section7) ──
    feas_data = vl_report.get("section7_execution_feasibility", {})
    if feas_data:
        p = doc.add_paragraph()
        run = p.add_run("── 执行可行性 ──")
        run.bold = True
        run.font.size = Pt(11)

        proto_days = feas_data.get("prototype_days", "?")
        process_diff = feas_data.get("process_difficulty", "?")
        supply_risk = feas_data.get("supply_chain_risk", "?")
        is_new_cat = feas_data.get("is_new_category", False)
        new_processes = feas_data.get("new_process_needed", [])
        feas_summary = feas_data.get("feasibility_summary", "")

        feas_items = [
            ("打样周期", f"{proto_days}天"),
            ("工艺难度", process_diff),
            ("供应链风险", supply_risk),
        ]
        if is_new_cat:
            feas_items.append(("新品类", "是（额外供应链风险）"))
        if new_processes:
            feas_items.append(("需新开工艺", ", ".join(str(p) for p in new_processes[:5])))
        _add_kv_table(doc, feas_items)

        if feas_summary:
            doc.add_paragraph(f"总结: {feas_summary}")

    # ── 下一步建议 (section9) ──
    next_steps = vl_report.get("section9_next_steps", {})
    if next_steps:
        suggestions = next_steps.get("suggestions", [])
        summary = next_steps.get("summary", "")
        if suggestions:
            _add_bullet_list(doc, "下一步建议", suggestions[:5], COLOR_BLUE, "💡")
        if summary:
            doc.add_paragraph(f"综合建议: {summary}")


def _add_module_comparison_table(doc, module_comparison: list[dict]):
    """
    将逐模块差距矩阵渲染为 Word 表格。
    列：模块名称 | 我方状态 | 竞品/参考 | 差距等级 | 用户感知 | 优先级 | 建议
    """
    if not module_comparison:
        return

    headers = ["模块", "我方状态", "竞品/参考", "差距", "感知", "优先级", "建议"]
    num_cols = len(headers)
    num_rows = len(module_comparison) + 1

    table = doc.add_table(rows=num_rows, cols=num_cols, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "微软雅黑"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 数据行
    for i, mc in enumerate(module_comparison):
        row = table.rows[i + 1]
        values = [
            mc.get("module_name", ""),
            mc.get("our_status", ""),
            mc.get("competitor_status", ""),
            mc.get("gap_level", ""),
            mc.get("user_perception", ""),
            mc.get("upgrade_priority", ""),
            mc.get("suggestion", ""),
        ]
        for j, val in enumerate(values):
            cell = row.cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "微软雅黑"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 差距列着色
            if j == 3 and val in GAP_COLORS:
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = GAP_COLORS[val]
                    run.bold = True

            # 优先级列着色
            if j == 5:
                priority_colors = {
                    "P0": COLOR_RED,
                    "P1": RGBColor(0xFF, 0x99, 0x00),
                    "P2": COLOR_BLUE,
                    "P3": COLOR_GRAY,
                }
                if val in priority_colors:
                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = priority_colors[val]
                        run.bold = True

    # 设置列宽
    try:
        widths = [Cm(2.0), Cm(2.5), Cm(2.5), Cm(1.2), Cm(1.2), Cm(1.2), Cm(4.5)]
        for row in table.rows:
            for j, width in enumerate(widths):
                row.cells[j].width = width
    except Exception:
        pass  # 列宽设置失败不影响输出


def _add_roadmap_table(doc, roadmap: list[dict]):
    """将升级路线图渲染为 Word 表格"""
    if not roadmap:
        return

    headers = ["优先级", "模块/方向", "具体行动", "预期效果"]
    num_cols = len(headers)
    num_rows = len(roadmap) + 1

    table = doc.add_table(rows=num_rows, cols=num_cols, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    # 数据行
    priority_colors = {
        "P0": COLOR_RED,
        "P1": RGBColor(0xFF, 0x99, 0x00),
        "P2": COLOR_BLUE,
        "P3": COLOR_GRAY,
    }
    for i, item in enumerate(roadmap):
        row = table.rows[i + 1]
        values = [
            item.get("priority", ""),
            item.get("module_name", item.get("direction", "")),
            item.get("action", item.get("suggestion", "")),
            item.get("expected_effect", item.get("reason", "")),
        ]
        for j, val in enumerate(values):
            cell = row.cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "微软雅黑"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 优先级列着色
            if j == 0 and val in priority_colors:
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = priority_colors[val]
                    run.bold = True


def _add_summary_section(doc, summary: dict):
    """添加分析摘要（优势/不足/复用/新建）"""
    sections = [
        ("我方优势", "our_strengths"),
        ("我方不足", "our_weaknesses"),
        ("可复用模块", "reuse_modules"),
        ("需新建模块", "new_modules_needed"),
    ]
    for label, key in sections:
        items = summary.get(key, [])
        if items:
            color = COLOR_GREEN if key in ("our_strengths", "reuse_modules") else COLOR_RED
            _add_bullet_list(doc, label, items, color)


def _add_bullet_list(doc, label: str, items: list, color: RGBColor, prefix: str = ""):
    """添加带标题的要点列表"""
    p = doc.add_paragraph()
    display_label = f"{prefix} {label}" if prefix else label
    run = p.add_run(display_label)
    run.bold = True
    run.font.color.rgb = color

    for item in items:
        bullet = doc.add_paragraph(str(item), style="List Bullet")
        for run in bullet.runs:
            run.font.size = Pt(10)


def _extract_module_comparison(specific_analysis: dict) -> list[dict]:
    """从专项分析结果中提取 module_comparison"""
    if not specific_analysis:
        return []

    # 1. 可能在顶层
    mc = specific_analysis.get("module_comparison", [])
    if mc:
        return mc

    # 2. 可能在 comparison 子对象中
    comparison = specific_analysis.get("comparison", {})
    if isinstance(comparison, dict):
        mc = comparison.get("module_comparison", [])
        if mc:
            return mc

        # 品类缺失V2: module_reuse_detail 格式转换
        mrd = comparison.get("module_reuse_detail", [])
        if mrd:
            rows = []
            for m in mrd:
                rows.append({
                    "module_name": m.get("module_name", ""),
                    "our_status": m.get("source", ""),
                    "competitor_status": m.get("note", ""),
                    "gap_level": {"相同": "无", "相似": "低", "需改造": "中", "需新建": "高"}.get(m.get("match_level", ""), "中"),
                    "user_perception": "",
                    "upgrade_priority": "",
                    "suggestion": f"{m.get('match_level', '')} ({m.get('module_type', '')})",
                })
            return rows

    # 3. 可能在 vl_report.section3_module_comparison（VL对比拆解结果）
    vl_report = specific_analysis.get("vl_report", {})
    if isinstance(vl_report, dict):
        section3 = vl_report.get("section3_module_comparison", {})
        if isinstance(section3, dict):
            # 转换section3为表格格式
            rows = []
            # 相同模块
            for m in section3.get("same_modules", []):
                rows.append({
                    "module_name": m.get("module_name", ""),
                    "our_status": m.get("own_detail", ""),
                    "competitor_status": m.get("competitor_detail", ""),
                    "gap_level": "无",
                    "user_perception": "",
                    "upgrade_priority": "",
                    "suggestion": "复用现有模块" if m.get("reuse") else "可复用",
                })
            # 竞品独有
            for m in section3.get("competitor_only", []):
                hit = m.get("upgrade_direction_hit", False)
                rows.append({
                    "module_name": m.get("module_name", ""),
                    "our_status": "缺失",
                    "competitor_status": m.get("detail", ""),
                    "gap_level": "高",
                    "user_perception": "",
                    "upgrade_priority": "P0" if hit else "P1",
                    "suggestion": "需新建或升级" + (" ← 升级方向命中" if hit else ""),
                })
            # 自家独有
            for m in section3.get("own_only", []):
                rows.append({
                    "module_name": m.get("module_name", ""),
                    "our_status": m.get("detail", ""),
                    "competitor_status": "缺失",
                    "gap_level": "无",
                    "user_perception": "",
                    "upgrade_priority": "",
                    "suggestion": "自家差异化优势" if m.get("is_advantage") else "保持",
                })
            # 结构差异
            for s in section3.get("structural_differences", []):
                rows.append({
                    "module_name": s.get("aspect", ""),
                    "our_status": s.get("own", ""),
                    "competitor_status": s.get("competitor", ""),
                    "gap_level": "中",
                    "user_perception": "",
                    "upgrade_priority": "P2",
                    "suggestion": "需评估工艺差异",
                })
            if rows:
                return rows

    return []


def _extract_roadmap(specific_analysis: dict) -> list[dict]:
    """从专项分析结果中提取 upgrade_roadmap"""
    if not specific_analysis:
        return []

    roadmap = specific_analysis.get("upgrade_roadmap", [])
    if roadmap:
        return roadmap

    comparison = specific_analysis.get("comparison", {})
    if isinstance(comparison, dict):
        roadmap = comparison.get("upgrade_roadmap", [])
        if roadmap:
            return roadmap

    # 从 vl_report 的 next_steps 和 competitor_only 提取
    vl_report = specific_analysis.get("vl_report", {})
    if isinstance(vl_report, dict):
        roadmap_items = []

        # 竞品独有模块 → P0/P1 行动项
        section3 = vl_report.get("section3_module_comparison", {})
        for m in section3.get("competitor_only", []):
            hit = m.get("upgrade_direction_hit", False)
            roadmap_items.append({
                "priority": "P0" if hit else "P1",
                "module_name": m.get("module_name", ""),
                "action": f"补齐竞品独有模块: {m.get('detail', '')[:60]}",
                "expected_effect": "缩小与竞品差距" if hit else "提升产品竞争力",
            })

        # next_steps.suggestions
        next_steps = vl_report.get("section9_next_steps", {})
        for s in next_steps.get("suggestions", [])[:3]:
            roadmap_items.append({
                "priority": "P2",
                "module_name": "综合建议",
                "action": s,
                "expected_effect": "",
            })

        if roadmap_items:
            return roadmap_items

    return []


def _extract_summary(specific_analysis: dict) -> dict:
    """从专项分析结果中提取 summary"""
    if not specific_analysis:
        return {}

    summary = specific_analysis.get("summary", {})
    if summary:
        return summary

    comparison = specific_analysis.get("comparison", {})
    if isinstance(comparison, dict):
        summary = comparison.get("summary", {})
        if summary:
            return summary

    # 从 vl_report 构建摘要
    vl_report = specific_analysis.get("vl_report", {})
    if isinstance(vl_report, dict):
        result = {}
        section3 = vl_report.get("section3_module_comparison", {})
        reuse_data = vl_report.get("section5_module_reuse", {})

        same = [m.get("module_name", "") for m in section3.get("same_modules", [])]
        comp_only = [m.get("module_name", "") for m in section3.get("competitor_only", [])]
        own_only = [m.get("module_name", "") for m in section3.get("own_only", [])]
        new_mods = reuse_data.get("new_modules_needed", [])

        if own_only:
            result["our_strengths"] = own_only
        if comp_only:
            result["our_weaknesses"] = comp_only
        if same:
            result["reuse_modules"] = same
        if new_mods:
            result["new_modules_needed"] = [str(m) for m in new_mods]

        if result:
            return result

    return {}


def _extract_design_requirements(report_text: str) -> str:
    """从报告文本中提取设计要求段落"""
    if not report_text:
        return ""

    lines = report_text.split("\n")
    capturing = False
    captured = []

    for line in lines:
        stripped = line.strip()
        if "附" in stripped and "设计" in stripped:
            capturing = True
            continue
        if capturing:
            if stripped.startswith(("一、", "二、", "三、", "四、", "五、", "六、")):
                break
            if stripped and not stripped.startswith("====") and not stripped.startswith("----"):
                captured.append(stripped)

    return "\n".join(captured)


def _add_hot_upgrade_dimension_details(doc, specific_score: dict):
    """为爆品升级添加5维度详细评语段落"""
    dimensions = specific_score.get("dimensions", [])
    if not dimensions:
        return

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("── 各维度详细评语 ──")
    run.bold = True
    run.font.size = Pt(11)

    dim_labels = {
        "模块复用": "模块复用（48分）: 核心/非核心模块复用 + 升级方向匹配 + 市场验证",
        "模块升级合理性": "模块升级合理性（22分）: 目标-动作一致性 + 卖点保留度 + 升级必要性",
        "价格分析": "价格分析（10分）: 价格定位合理性 + 成本-定价匹配度",
        "营销分析": "营销分析（10分）: 升级卖点的营销价值",
        "可行性分析": "可行性分析（10分）: 供应链/打样可行性",
    }

    for dim in dimensions:
        if not isinstance(dim, dict):
            continue
        name = dim.get("name", "")
        score = dim.get("score", 0)
        max_s = dim.get("max_score", 10)
        reason = dim.get("reason", "")

        label = dim_labels.get(name, f"{name}（{max_s}分）")
        p = doc.add_paragraph()
        run = p.add_run(f"【{label}】")
        run.bold = True
        run.font.size = Pt(10.5)

        score_color = COLOR_GREEN if score >= max_s * 0.7 else (RGBColor(0xFF, 0x99, 0x00) if score >= max_s * 0.4 else COLOR_RED)
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f"得分: {score}/{max_s}")
        run2.bold = True
        run2.font.color.rgb = score_color

        if reason:
            for line in reason.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)


def _get_specific_total(specific_score: dict) -> int:
    """从专项评分中获取总分"""
    if not specific_score:
        return 0
    if isinstance(specific_score, dict):
        return specific_score.get("total_score", 0)
    if hasattr(specific_score, "total_score"):
        return specific_score.total_score
    return 0


def _add_category_gap_dimension_details(doc, specific_score: dict, scenario: str):
    """为品类缺失添加维度详细评语段落"""
    dimensions = specific_score.get("dimensions", [])
    if not dimensions:
        return

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("── 各维度详细评语 ──")
    run.bold = True
    run.font.size = Pt(11)

    if scenario == "A":
        dim_labels = {
            "模块复用": "模块复用（48分）: CBB sub_type匹配 + VL对比覆盖率",
            "模块升级合理性": "模块升级合理性（22分）: 目标-动作一致性 + 卖点保留度 + 升级必要性",
            "价格分析": "价格分析（10分）: 定价竞争力 + 成本-定价匹配度",
            "营销分析": "营销分析（10分）: 升级卖点的营销价值",
            "可行性分析": "可行性分析（10分）: 供应链/打样可行性",
        }
    else:
        dim_labels = {
            "模块可行性": "模块可行性（50分）: 面料可行性15 + 版型可行性15 + 其它模块可获取性12 + 开模难度8",
            "产品设计合理性": "产品设计合理性（30分）: 设计目的明确性 + 设计方向合理性 + 差异化价值",
            "价格与市场": "价格与市场（20分）: 价格竞争力 + 市场验证",
        }

    for dim in dimensions:
        if not isinstance(dim, dict):
            continue
        name = dim.get("name", "")
        score = dim.get("score", 0)
        max_s = dim.get("max_score", 10)
        reason = dim.get("reason", "")

        label = dim_labels.get(name, f"{name}（{max_s}分）")
        p = doc.add_paragraph()
        run = p.add_run(f"【{label}】")
        run.bold = True
        run.font.size = Pt(10.5)

        score_color = COLOR_GREEN if score >= max_s * 0.7 else (RGBColor(0xFF, 0x99, 0x00) if score >= max_s * 0.4 else COLOR_RED)
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f"得分: {score}/{max_s}")
        run2.bold = True
        run2.font.color.rgb = score_color

        if reason:
            for line in reason.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)


def _add_competitor_upgrade_dimension_details(doc, specific_score: dict):
    """为竞品升级添加5维度详细评语段落"""
    dimensions = specific_score.get("dimensions", [])
    if not dimensions:
        return

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("── 各维度详细评语 ──")
    run.bold = True
    run.font.size = Pt(11)

    dim_labels = {
        "模块复用基础": "模块复用基础（35分）: CBB模块匹配 + 自家-竞品模块重叠度",
        "卖点复制可行性": "卖点复制可行性（25分）: 竞品核心卖点可复制性 + 卖点保留度",
        "差异化超越空间": "差异化超越空间（20分）: 竞品优势可超越性 + 差异化方向",
        "价格竞争力": "价格竞争力（10分）: 升级后定价vs竞品价格偏差",
        "市场验证": "市场验证（10分）: 竞品销量 + 品牌数验证需求存在性",
    }

    for dim in dimensions:
        if not isinstance(dim, dict):
            continue
        name = dim.get("name", "")
        score = dim.get("score", 0)
        max_s = dim.get("max_score", 10)
        reason = dim.get("reason", "")

        label = dim_labels.get(name, f"{name}（{max_s}分）")
        p = doc.add_paragraph()
        run = p.add_run(f"【{label}】")
        run.bold = True
        run.font.size = Pt(10.5)

        score_color = COLOR_GREEN if score >= max_s * 0.7 else (RGBColor(0xFF, 0x99, 0x00) if score >= max_s * 0.4 else COLOR_RED)
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f"得分: {score}/{max_s}")
        run2.bold = True
        run2.font.color.rgb = score_color

        if reason:
            for line in reason.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)


def _add_low_sale_iterate_dimension_details(doc, specific_score: dict):
    """为未起量迭代添加5维度详细评语段落"""
    dimensions = specific_score.get("dimensions", [])
    if not dimensions:
        return

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("── 各维度详细评语 ──")
    run.bold = True
    run.font.size = Pt(11)

    dim_labels = {
        "问题诊断": "问题诊断（25分）: 没卖好的原因是否找对",
        "迭代-诊断一致性": "迭代-诊断一致性（25分）: 提出的改动能否解决诊断问题",
        "模块复用基础": "模块复用基础（20分）: 现有CBB模块复用率",
        "增量空间": "增量空间（15分）: 竞品验证 + 市场容量",
        "风险可控度": "风险可控度（15分）: 已起量产品冲突 + 新建成本",
    }

    for dim in dimensions:
        if not isinstance(dim, dict):
            continue
        name = dim.get("name", "")
        score = dim.get("score", 0)
        max_s = dim.get("max_score", 10)
        reason = dim.get("reason", "")

        label = dim_labels.get(name, f"{name}（{max_s}分）")
        p = doc.add_paragraph()
        run = p.add_run(f"【{label}】")
        run.bold = True
        run.font.size = Pt(10.5)

        score_color = COLOR_GREEN if score >= max_s * 0.7 else (RGBColor(0xFF, 0x99, 0x00) if score >= max_s * 0.4 else COLOR_RED)
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f"得分: {score}/{max_s}")
        run2.bold = True
        run2.font.color.rgb = score_color

        if reason:
            for line in reason.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)


def _add_category_gap_scenario_section(doc, specific_analysis: dict):
    """品类缺失场景判断信息"""
    gap_info = specific_analysis.get("gap_info", {})
    gap_type = specific_analysis.get("gap_type", "")
    if not gap_info:
        return

    gap_labels = {
        "brand_gap": "品牌缺失（公司有该品类产品，但立项品牌下没有）",
        "category_gap": "品类缺失（公司完全没有该品类产品）",
        "no_gap": "品类补全（同品牌下已有产品）",
    }

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("── 场景判断 ──")
    run.bold = True
    run.font.size = Pt(11)

    items = []
    label = gap_labels.get(gap_type, gap_type)
    items.append(("场景类型", label))
    if gap_info.get("gap_description"):
        items.append(("缺失描述", gap_info["gap_description"]))
    if items:
        _add_kv_table(doc, items)


def _add_cbb_match_section(doc, specific_analysis: dict):
    """CBB模块匹配详情"""
    cbb_match = specific_analysis.get("cbb_match")
    if not cbb_match or not hasattr(cbb_match, "module_matches"):
        return
    if not cbb_match.module_matches:
        return

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("── CBB模块匹配（FAISS语义检索） ──")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(f"匹配率: {cbb_match.match_rate}% ({cbb_match.matched}/{cbb_match.total})")
    run.bold = True

    for mm in cbb_match.module_matches[:10]:
        status = "✓" if mm.matched else "✗"
        score_str = f" score={mm.score:.2f}" if mm.score else ""
        modules_info = ""
        if mm.cbb_modules:
            modules_info = " → " + ", ".join(m["cbb_name"] for m in mm.cbb_modules[:2])
        color = COLOR_GREEN if mm.matched else COLOR_RED
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"[{status}] ")
        run.font.color.rgb = color
        p.add_run(f"{mm.vl_module} → {mm.cbb_category}/{mm.cbb_sub_type} [{mm.match_level}{score_str}]{modules_info}")


def _add_market_overview_section(doc, specific_analysis: dict):
    """品类市场概况"""
    market = specific_analysis.get("market_overview", {})
    if not market:
        return

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("── 品类市场概况 ──")
    run.bold = True
    run.font.size = Pt(11)

    total_products = market.get("total_products", 0)
    total_sales = market.get("total_category_sales", 0)

    items = []
    items.append(("品类产品总数", f"{total_products}个"))
    items.append(("品类累计总销量", f"{total_sales:,}"))
    _add_kv_table(doc, items)

    brand_dist = market.get("brand_distribution", [])
    if brand_dist:
        p = doc.add_paragraph()
        run = p.add_run("品牌分布:")
        run.bold = True
        for bd in brand_dist[:5]:
            pct = (bd["total_sales"] / total_sales * 100) if total_sales > 0 else 0
            doc.add_paragraph(
                f"{bd['brand']}: {bd['product_count']}个产品, 销量{bd['total_sales']:,} ({pct:.1f}%)",
                style="List Bullet"
            )


def _add_single_vl_report(doc, vl_report: dict):
    """单拆模式的VL报告（场景B：无对比，只有竞品模块拆解）"""
    b_level = vl_report.get("section2_abc_modules", {}).get("b_level", [])
    if not b_level:
        return

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("── 竞品模块拆解 ──")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(f"共拆解出 {len(b_level)} 个模块")
    run.bold = True

    for m in b_level[:10]:
        name = m.get("name", "?")
        func = m.get("core_function", "")
        material = m.get("typical_material", "")
        priority = m.get("priority", "")
        parts = [name]
        if func:
            parts.append(func)
        if material:
            parts.append(f"材料:{material}")
        if priority:
            parts.append(f"优先级:{priority}")
        doc.add_paragraph(" | ".join(parts), style="List Bullet")

    # 视觉分析
    section1 = vl_report.get("section1_visual_analysis", {})
    if section1:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("── 视觉分析 ──")
        run.bold = True
        run.font.size = Pt(11)

        vis_items = []
        if section1.get("product_type"):
            vis_items.append(("产品类型", section1["product_type"]))
        if section1.get("structure_form"):
            vis_items.append(("结构形态", section1["structure_form"]))
        if section1.get("material_texture"):
            vis_items.append(("材料质感", section1["material_texture"]))
        if vis_items:
            _add_kv_table(doc, vis_items)
