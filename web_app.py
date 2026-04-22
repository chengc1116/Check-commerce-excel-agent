# -*- coding: utf-8 -*-
"""
产品立项审核系统 — Web UI 后端

FastAPI 服务，提供 REST API 供前端调用。
所有数据统一存储在 project_review.db。
启动方式: python web_app.py
"""

from __future__ import annotations

import asyncio
import hashlib
import json
import os
import secrets
import sqlite3
import sys
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional

# Windows UTF-8
if sys.platform == "win32":
    try:
        sys.stdout = open(sys.stdout.fileno(), mode="w", encoding="utf-8", buffering=1)
        sys.stderr = open(sys.stderr.fileno(), mode="w", encoding="utf-8", buffering=1)
    except Exception:
        pass

# 加载 .env 环境变量（必须在其他模块 import 之前）
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent / ".env", override=True)
except ImportError:
    pass

from fastapi import FastAPI, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

# ============================================================
# App 创建
# ============================================================

app = FastAPI(title="产品立项审核系统", version="2.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 静态文件
STATIC_DIR = Path(__file__).parent / "static"
STATIC_DIR.mkdir(exist_ok=True)
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")

# 项目根目录
PROJECT_ROOT = Path(__file__).parent.resolve()

# 上传目录
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", str(PROJECT_ROOT / "uploads"))).resolve()
UPLOAD_DIR.mkdir(exist_ok=True)

# 统一数据库（products + cbb_modules + reviews + users）
DB_PATH = PROJECT_ROOT / "data" / "project_review.db"

# 旧数据库（迁移完成后可手动删除）
OLD_REVIEW_DB = PROJECT_ROOT / "data" / "reviews.db"

# 每页条数
PAGE_SIZE = 10


# ============================================================
# DB 连接辅助
# ============================================================

def _get_db():
    """连接统一数据库 project_review.db，确保 reviews + users 表存在"""
    db_path = str(DB_PATH)
    os.makedirs(str(DB_PATH.parent), exist_ok=True)
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("""
        CREATE TABLE IF NOT EXISTS reviews (
            id TEXT PRIMARY KEY,
            file_name TEXT NOT NULL,
            task_type TEXT NOT NULL,
            task_label TEXT DEFAULT '',
            overall_score INTEGER DEFAULT 0,
            risk_level TEXT DEFAULT '未知',
            report TEXT DEFAULT '',
            project_data TEXT DEFAULT '{}',
            specific_score TEXT DEFAULT '{}',
            common_scores TEXT DEFAULT '{}',
            elapsed_seconds REAL DEFAULT 0,
            status TEXT DEFAULT 'processing',
            error TEXT DEFAULT '',
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            role TEXT NOT NULL DEFAULT 'user',
            display_name TEXT DEFAULT '',
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    return conn


def _product_row_to_dict(row: sqlite3.Row) -> dict:
    """将 products 表的行映射为前端统一格式"""
    d = dict(row)
    return {
        "id": d.get("rowid") or d.get("id") or 0,
        "sku": d.get("product_code") or d.get("sku", ""),
        "brand": d.get("brand", ""),
        "category_l1": d.get("category1") or d.get("category_l1", ""),
        "category_l2": d.get("category2") or d.get("category_l2", ""),
        "category_l3": d.get("category3") or d.get("category_l3", ""),
        "version": d.get("version", ""),
        "status": (d.get("status") or "active").lower(),
        "image_path": d.get("image_url") or d.get("image_path", ""),
        "remark": d.get("remark", ""),
        "created_at": d.get("created_at", ""),
        "updated_at": d.get("updated_at", ""),
    }


# ============================================================
# 用户认证系统
# ============================================================

def _hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def _init_users_db():
    """初始化默认账号"""
    conn = _get_db()
    default_users = [
        ("admin", "admin123", "super_admin", "超级管理员"),
        ("zhangsan", "zhang123", "user", "张三"),
        ("lisi", "li123", "user", "李四"),
        ("wangwu", "wang123", "user", "王五"),
        ("admin1", "123456", "user", "Admin1"),
    ]
    for uname, pwd, role, dname in default_users:
        existing = conn.execute("SELECT id FROM users WHERE username = ?", (uname,)).fetchone()
        if not existing:
            conn.execute(
                "INSERT INTO users (username, password_hash, role, display_name) VALUES (?,?,?,?)",
                (uname, _hash_password(pwd), role, dname),
            )
    conn.commit()
    conn.close()


def _migrate_old_db():
    """从旧 reviews.db 迁移 reviews + users 到 project_review.db"""
    if not OLD_REVIEW_DB.exists():
        return
    try:
        old_conn = sqlite3.connect(str(OLD_REVIEW_DB))
        old_conn.row_factory = sqlite3.Row

        # 迁移 reviews
        try:
            old_reviews = old_conn.execute("SELECT * FROM reviews").fetchall()
            if old_reviews:
                new_conn = _get_db()
                for r in old_reviews:
                    # 检查是否已存在（避免重复迁移）
                    exists = new_conn.execute("SELECT id FROM reviews WHERE id=?", (r["id"],)).fetchone()
                    if not exists:
                        new_conn.execute(
                            """INSERT INTO reviews (id,file_name,task_type,task_label,overall_score,
                               risk_level,report,project_data,specific_score,common_scores,
                               elapsed_seconds,status,error,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                            (r["id"], r["file_name"], r["task_type"], r["task_label"], r["overall_score"],
                             r["risk_level"], r["report"], r["project_data"], r["specific_score"],
                             r["common_scores"], r["elapsed_seconds"], r["status"], r["error"], r["created_at"]),
                        )
                new_conn.commit()
                new_conn.close()
                print(f"[Migration] 迁移 {len(old_reviews)} 条审核记录 → project_review.db")
        except Exception as e:
            print(f"[Migration] reviews 迁移失败: {e}")

        # 迁移 users（仅迁移旧 DB 中可能有而新 DB 没有的用户）
        try:
            old_users = old_conn.execute("SELECT * FROM users").fetchall()
            if old_users:
                new_conn = _get_db()
                migrated = 0
                for u in old_users:
                    exists = new_conn.execute("SELECT id FROM users WHERE username=?", (u["username"],)).fetchone()
                    if not exists:
                        new_conn.execute(
                            "INSERT INTO users (username,password_hash,role,display_name,created_at) VALUES (?,?,?,?,?)",
                            (u["username"], u["password_hash"], u["role"], u["display_name"], u["created_at"]),
                        )
                        migrated += 1
                new_conn.commit()
                new_conn.close()
                if migrated:
                    print(f"[Migration] 迁移 {migrated} 个用户 → project_review.db")
        except Exception as e:
            print(f"[Migration] users 迁移失败: {e}")

        old_conn.close()
        # 迁移成功后重命名旧文件（不直接删，保险）
        backup_path = str(OLD_REVIEW_DB) + ".bak"
        if not Path(backup_path).exists():
            os.rename(str(OLD_REVIEW_DB), backup_path)
            print(f"[Migration] 旧 reviews.db → {backup_path}")
    except Exception as e:
        print(f"[Migration] 迁移过程出错: {e}")


# 启动时初始化
_init_users_db()
_migrate_old_db()

# Token 存储
_tokens: dict[str, dict] = {}


def _create_token(username: str, role: str, display_name: str) -> str:
    token = secrets.token_hex(32)
    _tokens[token] = {
        "username": username,
        "role": role,
        "display_name": display_name,
        "expires_at": datetime.now() + timedelta(hours=24),
    }
    return token


def _get_current_user(request: Request) -> Optional[dict]:
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        token = auth[7:]
        info = _tokens.get(token)
        if info and info["expires_at"] > datetime.now():
            return info
    return None


def _require_auth(request: Request) -> dict:
    user = _get_current_user(request)
    if not user:
        raise HTTPException(401, "未登录或登录已过期")
    return user


def _require_super_admin(request: Request) -> dict:
    user = _require_auth(request)
    if user["role"] != "super_admin":
        raise HTTPException(403, "权限不足，仅超级管理员可执行此操作")
    return user


# ============================================================
# API 路由 — 认证
# ============================================================

class LoginRequest(BaseModel):
    username: str
    password: str


@app.post("/api/auth/login")
async def login(req: LoginRequest):
    conn = _get_db()
    user = conn.execute("SELECT * FROM users WHERE username = ?", (req.username,)).fetchone()
    conn.close()
    if not user or user["password_hash"] != _hash_password(req.password):
        raise HTTPException(401, "用户名或密码错误")
    token = _create_token(user["username"], user["role"], user["display_name"])
    return {
        "token": token,
        "user": {
            "username": user["username"],
            "role": user["role"],
            "display_name": user["display_name"],
        },
    }


@app.get("/api/auth/me")
async def get_me(request: Request):
    user = _get_current_user(request)
    if not user:
        raise HTTPException(401, "未登录")
    return {"username": user["username"], "role": user["role"], "display_name": user["display_name"]}


@app.post("/api/auth/logout")
async def logout(request: Request):
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        token = auth[7:]
        _tokens.pop(token, None)
    return {"ok": True}


# ============================================================
# API 路由 — 仪表盘
# ============================================================

@app.get("/")
async def index():
    return FileResponse(str(STATIC_DIR / "index.html"))


@app.get("/api/dashboard")
async def dashboard():
    try:
        conn = _get_db()
        total = conn.execute("SELECT COUNT(*) FROM products").fetchone()[0]
        active = conn.execute("SELECT COUNT(*) FROM products WHERE status = 'ACTIVE'").fetchone()[0]
        archived = conn.execute("SELECT COUNT(*) FROM products WHERE status != 'ACTIVE'").fetchone()[0]
        cats = conn.execute(
            "SELECT category2, COUNT(*) as cnt FROM products WHERE status = 'ACTIVE' GROUP BY category2 ORDER BY cnt DESC"
        ).fetchall()
        conn.close()
        stats = {
            "total": total,
            "active": active,
            "archived": archived,
            "sales_months": 0,
            "categories": [{"name": r["category2"], "count": r["cnt"]} for r in cats if r["category2"]],
        }
    except Exception as e:
        print(f"[Dashboard] 产品库查询失败: {e}")
        stats = {"total": 0, "active": 0, "archived": 0, "sales_months": 0, "categories": []}

    try:
        conn = _get_db()
        review_total = conn.execute("SELECT COUNT(*) FROM reviews").fetchone()[0]
        review_recent = conn.execute(
            "SELECT COUNT(*) FROM reviews WHERE created_at > datetime('now', '-7 days')"
        ).fetchone()[0]
        avg_row = conn.execute("SELECT AVG(overall_score) FROM reviews WHERE overall_score > 0").fetchone()
        avg_score = round(avg_row[0], 1) if avg_row and avg_row[0] else 0
        conn.close()
    except Exception:
        review_total = 0
        review_recent = 0
        avg_score = 0

    return {
        "product_stats": stats,
        "review_stats": {"total": review_total, "recent_7d": review_recent, "avg_score": avg_score},
    }


# ============================================================
# API 路由 — 项目审核
# ============================================================

@app.post("/api/review/start")
async def start_review(file: UploadFile = File(...), task_type: str = Form("hot_upgrade")):
    try:
        saved_name = f"{str(uuid.uuid4())[:8]}_{file.filename}"
        saved_path = UPLOAD_DIR / saved_name
        with open(saved_path, "wb") as f:
            f.write(await file.read())

        task_labels = {
            "hot_upgrade": "🔥 爆品升级",
            "competitor_upgrade": "⚔️ 竞品升级",
            "low_sale_iterate": "📉 未起量迭代",
            "category_gap": "🗺️ 品类地图缺失",
        }
        review_id = uuid.uuid4().hex[:12]
        conn = _get_db()
        conn.execute(
            "INSERT INTO reviews (id, file_name, task_type, task_label, status) VALUES (?,?,?,?,?)",
            (review_id, file.filename, task_type, task_labels.get(task_type, task_type), "processing"),
        )
        conn.commit()
        conn.close()

        asyncio.create_task(_run_review(review_id, str(saved_path), task_type))
        return {"review_id": review_id, "file_name": file.filename, "task_type": task_type, "task_label": task_labels.get(task_type, task_type)}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"审核启动失败: {str(e)}")


async def _run_review(review_id: str, file_path: str, task_type: str):
    print(f"[_run_review] 开始, review_id={review_id}, file={file_path}, task={task_type}")
    try:
        from product_review_agent.pipeline import run_pipeline
        result = await run_pipeline(file_path, task_type)
        print(f"[_run_review] pipeline 完成, score={result.overall_score}, risk={result.risk_level}")

        specific_score_dict = {}
        if result.specific_score:
            if hasattr(result.specific_score, "to_dict"):
                specific_score_dict = result.specific_score.to_dict()
            elif isinstance(result.specific_score, dict):
                specific_score_dict = result.specific_score

        try:
            common_scores_json = json.dumps(result.common_scores, ensure_ascii=False)
        except (TypeError, ValueError):
            common_scores_json = "{}"
        try:
            specific_score_json = json.dumps(specific_score_dict, ensure_ascii=False)
        except (TypeError, ValueError):
            specific_score_json = "{}"
        try:
            project_data_json = json.dumps(result.project_data, ensure_ascii=False)
        except (TypeError, ValueError):
            project_data_json = "{}"

        conn = _get_db()
        conn.execute("""
            UPDATE reviews SET overall_score=?, risk_level=?, report=?,
                project_data=?, specific_score=?, common_scores=?,
                elapsed_seconds=?, status=?, error=? WHERE id=?
        """, (
            result.overall_score, result.risk_level, result.report,
            project_data_json, specific_score_json, common_scores_json,
            result.elapsed_seconds, "completed" if not result.error else "error",
            result.error or "", review_id,
        ))
        conn.commit()
        conn.close()
        print(f"[_run_review] 数据库更新成功, review_id={review_id}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"[_run_review] 异常: {e}")
        try:
            conn = _get_db()
            conn.execute("UPDATE reviews SET status=?, error=? WHERE id=?", ("error", str(e), review_id))
            conn.commit()
            conn.close()
        except Exception as e2:
            print(f"[_run_review] 写入错误到DB也失败: {e2}")


@app.get("/api/review/status/{review_id}")
async def review_status(review_id: str):
    conn = _get_db()
    row = conn.execute("SELECT id, file_name, task_type, task_label, overall_score, risk_level, elapsed_seconds, status, error, created_at FROM reviews WHERE id=?", (review_id,)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(404, "审核记录不存在")
    return dict(row)


@app.get("/api/review/result/{review_id}")
async def review_result(review_id: str):
    conn = _get_db()
    row = conn.execute("SELECT * FROM reviews WHERE id=?", (review_id,)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(404, "审核记录不存在")
    return dict(row)


@app.get("/api/reviews")
async def list_reviews(
    status: Optional[str] = None, task_type: Optional[str] = None,
    page: int = Query(1, ge=1),
):
    conn = _get_db()
    conditions, params = [], []
    if status:
        conditions.append("status=?")
        params.append(status)
    if task_type:
        conditions.append("task_type=?")
        params.append(task_type)
    where = f"WHERE {' AND '.join(conditions)}" if conditions else ""

    total = conn.execute(f"SELECT COUNT(*) FROM reviews {where}", params).fetchone()[0]
    offset = (page - 1) * PAGE_SIZE
    rows = conn.execute(
        f"SELECT id, file_name, task_type, task_label, overall_score, risk_level, elapsed_seconds, status, error, created_at FROM reviews {where} ORDER BY created_at DESC LIMIT ? OFFSET ?",
        params + [PAGE_SIZE, offset],
    ).fetchall()
    conn.close()

    total_pages = (total + PAGE_SIZE - 1) // PAGE_SIZE if total > 0 else 1
    return {"total": total, "page": page, "page_size": PAGE_SIZE, "total_pages": total_pages, "items": [dict(r) for r in rows]}


@app.delete("/api/reviews/{review_id}")
async def delete_review(review_id: str, request: Request):
    _require_super_admin(request)
    conn = _get_db()
    row = conn.execute("SELECT id FROM reviews WHERE id=?", (review_id,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "审核记录不存在")
    conn.execute("DELETE FROM reviews WHERE id=?", (review_id,))
    conn.commit()
    conn.close()
    return {"ok": True, "message": "审核记录已删除"}


@app.get("/api/review/export/{review_id}")
async def export_review_docx(review_id: str):
    """将审核报告导出为 Word 文档 (.docx)"""
    conn = _get_db()
    row = conn.execute("SELECT * FROM reviews WHERE id=?", (review_id,)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(404, "审核记录不存在")
    if row["status"] != "completed":
        raise HTTPException(400, "审核尚未完成，无法导出")

    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    data = dict(row)
    report_text = data.get("report", "")
    project_data = json.loads(data.get("project_data", "{}")) if data.get("project_data") else {}
    specific_score = json.loads(data.get("specific_score", "{}")) if data.get("specific_score") else {}
    common_scores = json.loads(data.get("common_scores", "{}")) if data.get("common_scores") else {}

    doc = Document()

    # ---- 全局样式 ----
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3

    # ---- 标题 ----
    title = doc.add_heading("产品立项审核报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # ---- 产品概览表 ----
    doc.add_heading("产品概览", level=1)

    overview_items = [
        ("产品名称", project_data.get("project_name", "(未填写)")),
        ("品牌", project_data.get("brand", "(未填写)")),
        ("品类", " > ".join(filter(None, [
            project_data.get("categoryl1", ""),
            project_data.get("categoryl2", ""),
            project_data.get("categoryl3", ""),
        ])) or "(未填写)"),
        ("负责人", project_data.get("applicant", "(未填写)")),
        ("审核类型", data.get("task_label", "")),
        ("综合评分", f"{data.get('overall_score', 0)}/100"),
        ("风险等级", data.get("risk_level", "未知")),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]

    overview_table = doc.add_table(rows=len(overview_items), cols=2, style="Light Grid Accent 1")
    overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(overview_items):
        overview_table.rows[i].cells[0].text = label
        overview_table.rows[i].cells[1].text = str(value)
        for cell in overview_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "微软雅黑"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        # 标签列加粗
        for run in overview_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    # ---- 立项信息 ----
    doc.add_heading("立项信息", level=1)
    time_items = [
        ("立项时间", project_data.get("project_time", "(未填写)")),
        ("设计时间", project_data.get("design_time", "(未填写)")),
        ("打样时间", project_data.get("proofing_time", "(未填写)")),
        ("上架时间", project_data.get("launch_time", "(未填写)")),
    ]
    time_table = doc.add_table(rows=len(time_items), cols=2, style="Light Grid Accent 1")
    for i, (label, value) in enumerate(time_items):
        time_table.rows[i].cells[0].text = label
        time_table.rows[i].cells[1].text = str(value)
        for cell in time_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "微软雅黑"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        for run in time_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    # ---- 市场与定价 ----
    doc.add_heading("市场与定价", level=1)
    market_items = [
        ("市场规模", project_data.get("market_size", "(未填写)")),
        ("目标销量", project_data.get("estimated_sales", "(未填写)")),
        ("定价", project_data.get("pricing", "(未填写)")),
        ("毛利率", project_data.get("gfm", "(未填写)")),
        ("ERP成本", project_data.get("ERP_price", "(未填写)")),
    ]
    market_table = doc.add_table(rows=len(market_items), cols=2, style="Light Grid Accent 1")
    for i, (label, value) in enumerate(market_items):
        market_table.rows[i].cells[0].text = label
        market_table.rows[i].cells[1].text = str(value)
        for cell in market_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "微软雅黑"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        for run in market_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    # ---- 人群分析 ----
    doc.add_heading("人群分析", level=1)
    audience = common_scores.get("audience", {})
    if audience and audience.get("total_score"):
        _add_score_section(doc, audience, "人群")
    else:
        doc.add_paragraph("(暂无人群评分数据)")

    # ---- 场景分析 ----
    doc.add_heading("场景分析", level=1)
    scenario = common_scores.get("scenario", {})
    if scenario and scenario.get("total_score"):
        _add_score_section(doc, scenario, "场景")
    else:
        doc.add_paragraph("(暂无场景评分数据)")

    # ---- 专项分析 ----
    doc.add_heading(f"{data.get('task_label', '专项')}分析", level=1)
    if specific_score:
        _add_score_section(doc, specific_score, "专项")
    else:
        doc.add_paragraph("(暂无专项评分数据)")

    # ---- 完整报告文本 ----
    if report_text:
        doc.add_heading("完整审核报告", level=1)
        for line in report_text.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue
            # 检测各级标题（根据 pipeline 中的分隔线模式）
            if line.startswith("====") or line.startswith("----"):
                continue
            if line.startswith("一、") or line.startswith("二、") or line.startswith("三、") or \
               line.startswith("四、") or line.startswith("五、") or line.startswith("六、") or \
               line.startswith("七、") or line.startswith("附、"):
                h = doc.add_heading(line, level=2)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            elif line.startswith("  "):
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.left_indent = Cm(0.5)
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                p.paragraph_format.left_indent = Cm(0.5)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
            elif line.startswith("[") and "]" in line:
                # 小标题行，如 [评分明细]、[优势]、[不足] 等
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(10.5)
            else:
                doc.add_paragraph(line)

    # ---- 保存到临时文件 ----
    import tempfile
    tmp_dir = tempfile.mkdtemp()
    file_name = data.get("file_name", "report").replace(".xlsx", "").replace(".xls", "")
    docx_path = os.path.join(tmp_dir, f"审核报告_{file_name}.docx")
    doc.save(docx_path)

    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"审核报告_{file_name}.docx",
        background=lambda: _cleanup_temp(tmp_dir),
    )


def _add_score_section(doc, score_data: dict, section_name: str):
    """向文档添加一个评分板块（评分明细+优势+不足+建议）"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    total = score_data.get("total_score", 0)
    p = doc.add_paragraph()
    run = p.add_run(f"综合评分: {total}/100")
    run.bold = True
    run.font.size = Pt(11)

    # 评分维度表
    dimensions = score_data.get("dimensions", {})
    if dimensions:
        dim_items = [(name, info) for name, info in dimensions.items()
                     if isinstance(info, dict)]
        if dim_items:
            t = doc.add_table(rows=len(dim_items) + 1, cols=3, style="Light Grid Accent 1")
            # 表头
            for j, header in enumerate(["维度", "得分", "评价"]):
                t.rows[0].cells[j].text = header
                for run in t.rows[0].cells[j].paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(9)
            # 数据行
            for i, (name, info) in enumerate(dim_items):
                t.rows[i + 1].cells[0].text = name
                t.rows[i + 1].cells[1].text = f"{info.get('score', 0)}/{info.get('max_score', 25)}"
                t.rows[i + 1].cells[2].text = info.get("reason", "")
                for cell in t.rows[i + 1].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.name = "微软雅黑"
                            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 优势/不足/建议
    for label, key, color in [
        ("✅ 优势", "strengths", RGBColor(0x00, 0x80, 0x00)),
        ("⚠️ 不足", "weaknesses", RGBColor(0xCC, 0x33, 0x00)),
        ("💡 改进建议", "suggestions", RGBColor(0x00, 0x78, 0xD4)),
    ]:
        items = score_data.get(key, [])
        if items:
            p = doc.add_paragraph()
            run = p.add_run(label)
            run.bold = True
            run.font.color.rgb = color
            for item in items:
                bullet = doc.add_paragraph(item, style="List Bullet")
                for run in bullet.runs:
                    run.font.size = Pt(10)


def _cleanup_temp(tmp_dir: str):
    """清理临时目录"""
    import shutil
    try:
        shutil.rmtree(tmp_dir, ignore_errors=True)
    except Exception:
        pass


# ============================================================
# API 路由 — 产品库
# ============================================================

@app.get("/api/products/stats")
async def product_stats():
    try:
        conn = _get_db()
        total = conn.execute("SELECT COUNT(*) FROM products").fetchone()[0]
        active = conn.execute("SELECT COUNT(*) FROM products WHERE status = 'ACTIVE'").fetchone()[0]
        archived = total - active
        conn.close()
        return {"total": total, "active": active, "archived": archived}
    except Exception as e:
        raise HTTPException(500, str(e))


@app.get("/api/products")
async def list_products(
    status: Optional[str] = None, brand: Optional[str] = None,
    category_l2: Optional[str] = None, page: int = Query(1, ge=1),
):
    try:
        conn = _get_db()
        conditions, params = [], []
        if status:
            conditions.append("status = ?")
            params.append(status.upper())
        if brand:
            conditions.append("brand = ?")
            params.append(brand)
        if category_l2:
            conditions.append("category2 = ?")
            params.append(category_l2)
        where = f"WHERE {' AND '.join(conditions)}" if conditions else ""

        total = conn.execute(f"SELECT COUNT(*) FROM products {where}", params).fetchone()[0]
        offset = (page - 1) * PAGE_SIZE
        rows = conn.execute(f"SELECT *, rowid FROM products {where} ORDER BY updated_at DESC LIMIT ? OFFSET ?", params + [PAGE_SIZE, offset]).fetchall()
        conn.close()

        total_pages = (total + PAGE_SIZE - 1) // PAGE_SIZE if total > 0 else 1
        return {"total": total, "page": page, "page_size": PAGE_SIZE, "total_pages": total_pages, "items": [_product_row_to_dict(r) for r in rows]}
    except Exception as e:
        raise HTTPException(500, str(e))


@app.get("/api/products/{product_code}")
async def get_product(product_code: str, brand: Optional[str] = Query(None)):
    try:
        conn = _get_db()
        if brand:
            row = conn.execute("SELECT *, rowid FROM products WHERE product_code = ? AND brand = ?", (product_code, brand)).fetchone()
        else:
            row = conn.execute("SELECT *, rowid FROM products WHERE product_code = ? LIMIT 1", (product_code,)).fetchone()
        conn.close()
        if not row:
            raise HTTPException(404, "产品不存在")
        return _product_row_to_dict(row)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/products/import-inventory")
async def import_inventory(
    request: Request, file: UploadFile = File(...),
    sheet_name: Optional[str] = Form(None), month: Optional[str] = Form(None),
):
    _require_super_admin(request)
    saved_path = UPLOAD_DIR / f"inv_{str(uuid.uuid4())[:8]}_{file.filename}"
    with open(saved_path, "wb") as f:
        f.write(await file.read())

    try:
        from product_review_agent.product_db.inventory_parser import InventoryParser
        from product_review_agent.product_db.database import ProductDB

        parser = InventoryParser()
        products = parser.parse(str(saved_path), sheet_name=sheet_name, month=month, extract_images=True)
        db = ProductDB()
        for p in products:
            db.insert_product(p.to_dict())
            if p.sales_volume > 0 and p.month:
                db.upsert_sales(p.sku, p.month, p.sales_volume, p.brand)
        stats = db.get_stats()
        db.close()
        try:
            os.remove(saved_path)
        except Exception:
            pass
        return {"parsed_count": len(products), "stats": stats, "message": f"成功导入 {len(products)} 个产品"}
    except Exception as e:
        raise HTTPException(500, str(e))


# ============================================================
# API 路由 — CBB模块库
# ============================================================

@app.get("/api/cbb/stats")
async def cbb_stats():
    try:
        conn = _get_db()
        total = conn.execute("SELECT COUNT(*) FROM cbb_modules").fetchone()[0]
        categories = conn.execute("SELECT category, COUNT(*) as cnt FROM cbb_modules GROUP BY category ORDER BY cnt DESC").fetchall()
        conn.close()
        return {"total": total, "categories": [dict(r) for r in categories]}
    except Exception as e:
        return {"total": 0, "categories": [], "error": str(e)}


@app.get("/api/cbb")
async def list_cbb(
    category: Optional[str] = None, search: Optional[str] = None, page: int = Query(1, ge=1),
):
    try:
        conn = _get_db()
        conditions, params = [], []
        if category:
            conditions.append("category=?")
            params.append(category)
        if search:
            conditions.append("(cbb_code LIKE ? OR cbb_name LIKE ?)")
            params.extend([f"%{search}%", f"%{search}%"])
        where = f"WHERE {' AND '.join(conditions)}" if conditions else ""

        total = conn.execute(f"SELECT COUNT(*) FROM cbb_modules {where}", params).fetchone()[0]
        offset = (page - 1) * PAGE_SIZE
        rows = conn.execute(f"SELECT * FROM cbb_modules {where} ORDER BY usage_count DESC LIMIT ? OFFSET ?", params + [PAGE_SIZE, offset]).fetchall()
        conn.close()

        total_pages = (total + PAGE_SIZE - 1) // PAGE_SIZE if total > 0 else 1
        return {"total": total, "page": page, "page_size": PAGE_SIZE, "total_pages": total_pages, "items": [dict(r) for r in rows]}
    except Exception as e:
        raise HTTPException(500, str(e))


# ============================================================
# API 路由 — LLM状态
# ============================================================

@app.get("/api/llm/status")
async def llm_status():
    try:
        from product_review_agent.agents.llm_client import get_llm_client
        llm = get_llm_client()
        return {"available": llm.is_available, "model": llm.model, "vl_model": llm.vl_model, "base_url": llm.base_url}
    except Exception as e:
        return {"available": False, "error": str(e)}


# ============================================================
# 启动
# ============================================================

if __name__ == "__main__":
    import uvicorn
    host = os.getenv("HOST", "0.0.0.0")
    port = int(os.getenv("WEB_PORT", "8080"))
    print(f"🚀 产品立项审核系统 Web UI: http://localhost:{port}")
    uvicorn.run(app, host=host, port=port, log_level="info")
